# coding: utf-8

# Imports
from django.contrib.auth.models import User
from django.db import models

class TTypeUtilisateur(models.Model) :

	# Attributs
	code_type_util = models.CharField(max_length = 1, primary_key = True)
	int_type_util = models.CharField(max_length = 255)

	class Meta :
		db_table = 't_type_utilisateur'
		verbose_name = verbose_name_plural = 'T_TYPE_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_int_type_util(self) : return self.int_type_util
	def get_gpe_type_abs_set(self) : return self.tgroupetypeabsence_set
	def get_role_util_set(self) : return self.trolesutilisateur_set
	def get_util_set(self) : return self.tutilisateur_set

	def __str__(self) : return self.get_int_type_util()

class TMessage(models.Model) :

	# Attributs
	id_mess = models.AutoField(primary_key = True)
	corps_mess = models.TextField(blank = True, null = True)
	dt_mess = models.DateTimeField()
	obj_mess = models.CharField(blank = True, max_length = 255, null = True)

	class Meta :
		db_table = 't_message'
		ordering = ['-dt_mess']
		verbose_name = verbose_name_plural = 'T_MESSAGE'

	# Getters
	def get_pk(self) : return self.pk
	def get_corps_mess(self) : return self.corps_mess
	def get_corps_mess__text(self) :
		from bs4 import BeautifulSoup; return BeautifulSoup(self.get_corps_mess(), 'html.parser').get_text()
	def get_dt_mess(self) : return self.dt_mess
	def get_dt_mess__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_dt_mess())
	def get_obj_mess(self) : return self.obj_mess or 'Sans objet'
	def get_mess_util_set(self) : return self.tmessagesutilisateur_set
	def get_util_set(self) : return self.tutilisateur_set

	def __str__(self) :
		return '{0} - {1} - {2}'.format(self.get_emett_mess(), self.get_dt_mess__str(), self.get_obj_mess())

	def get_emett_mess(self) : from app.apps import AppConfig; return AppConfig.verbose_name

class TAnnee(models.Model) :

	# Import
	from django.contrib.postgres.fields import ArrayField

	# Attributs
	num_annee = models.IntegerField(primary_key = True, verbose_name = 'Année')
	plage_conges_annee = ArrayField(base_field = models.DateField(), size = 2)
	plage_rtt_annee = ArrayField(base_field = models.DateField(), size = 2)

	class Meta :
		db_table = 't_annee'
		ordering = ['-pk']
		verbose_name = verbose_name_plural = 'T_ANNEE'

	# Getters
	def get_pk(self) : return self.pk
	def get_plage_conges_annee(self) : return self.plage_conges_annee
	def get_plage_conges_annee__str(self) :
		from app.functions import get_local_format;
		return [get_local_format(elem) for elem in self.get_plage_conges_annee()]
	def get_plage_rtt_annee(self) : return self.plage_rtt_annee
	def get_plage_rtt_annee__str(self) :
		from app.functions import get_local_format;
		return [get_local_format(elem) for elem in self.get_plage_rtt_annee()]
	def get_abs_set(self) : return self.tabsence_set
	def get_bm_util_set(self) : return self.tbonusmalusutilisateur_set
	def get_decompt_util_set(self) : return self.tdecomptesutilisateur_set
	def get_statut_util_set(self) : return self.tstatutsutilisateur_set
	def get_trans_cet_util(self) : return self.ttransactionscetutilisateur_set

	def __str__(self) : return str(self.get_pk())

	@classmethod
	def create(_class, _annee) :

		# Import
		from datetime import date

		return _class.objects.create(
			pk = _annee,
			plage_conges_annee = [date(_annee, 1, 1), date(_annee + 1, 5, 31)],
			plage_rtt_annee = [date(_annee, 1, 1), date(_annee, 12, 31)]
		)

class TGroupeUtilisateur(models.Model) :

	# Import
	from app.validators import valid_cdc

	# Attributs
	id_gpe_util = models.AutoField(primary_key = True)
	int_gpe_util = models.CharField(max_length = 255, validators = [valid_cdc], verbose_name = 'Intitulé')

	class Meta :
		db_table = 't_groupe_utilisateur'
		ordering = ['int_gpe_util']
		verbose_name = verbose_name_plural = 'T_GROUPE_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_int_gpe_util(self) : return self.int_gpe_util
	def get_gpe_util_set(self) : return self.tgroupesutilisateur_set
	def get_util_set(self) : return self.tutilisateur_set

	def __str__(self) : return self.get_int_gpe_util()

	'''
	Précision sur le nombre d'utilisateurs composant un groupe
	self : Instance TGroupeUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_str_complet(self) :

		# Stockage du nombre d'utilisateurs en activité composant le groupe
		nbre_util_gpe = self.get_util_set().filter(en_act = True).count()

		# Mise en forme finale du libellé
		compl_str = '{} utilisateur'.format(nbre_util_gpe)
		if nbre_util_gpe > 1 : compl_str += 's'
		compl_str += ' en activité'

		return '{0} <span class="i">({1})</span>'.format(self, compl_str)

class TUtilisateur(User) :

	# Import
	from app.validators import valid_float

	# Attributs
	courr_second_util = models.EmailField(blank = True, null = True, verbose_name = 'Courriel secondaire')
	email_auto_courr_princ = models.BooleanField(
		default = False, verbose_name = 'Les notifications doivent-elles être envoyées sur le courriel principal ?'
	)
	email_auto_courr_second = models.BooleanField(
		default = False, verbose_name = 'Les notifications doivent-elles être envoyées sur le courriel secondaire ?'
	)
	en_act = models.BooleanField(default = True, verbose_name = 'L\'agent est-il en activité ?')
	est_super_secr = models.BooleanField(default = False)
	solde_cet_util = models.FloatField(
		default = 0, validators = [valid_float], verbose_name = 'Solde initial sur le compte épargne temps (CET)'
	)
	bm_util = models.ManyToManyField(TAnnee, related_name = '+', through = 'TBonusMalusUtilisateur')
	decompt_util = models.ManyToManyField(TAnnee, related_name = '+', through = 'TDecomptesUtilisateur')
	gpe_util = models.ManyToManyField(
		TGroupeUtilisateur,
		blank = True,
		through = 'TGroupesUtilisateur',
		verbose_name = 'Groupes d\'utilisateur|Intitulé'
	)
	mess = models.ManyToManyField(TMessage, through = 'TMessagesUtilisateur')
	statut_util = models.ManyToManyField(TAnnee, related_name = '+', through = 'TStatutsUtilisateur')
	trans_cet_util = models.ManyToManyField(TAnnee, related_name = '+', through = 'TTransactionsCetUtilisateur')
	type_util = models.ManyToManyField(
		TTypeUtilisateur,
		through = 'TRolesUtilisateur',
		verbose_name = 'Quels rôles doivent-ils être assignés au compte ?|Intitulé'
	)

	class Meta :
		db_table = 't_utilisateur'
		ordering = ['last_name', 'first_name']
		verbose_name = verbose_name_plural = 'T_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_courr_second_util(self) : return self.courr_second_util
	def get_date_joined(self) : return self.date_joined
	def get_date_joined__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_date_joined())
	def get_email(self) : return self.email
	def get_email_auto_courr_princ(self) : return self.email_auto_courr_princ
	def get_email_auto_courr_second(self) : return self.email_auto_courr_second
	def get_en_act(self) : return self.en_act
	def get_est_super_secr(self) : return self.est_super_secr
	def get_first_name(self) : return self.first_name
	def get_is_active(self) : return self.is_active
	def get_is_staff(self) : return self.is_staff
	def get_is_superuser(self) : return self.is_superuser
	def get_last_login(self) : return self.last_login
	def get_last_login__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_last_login())
	def get_last_name(self) : return self.last_name
	def get_solde_cet_util(self) : return self.solde_cet_util
	def get_solde_cet_util__str(self) : return '{0:g}'.format(self.get_solde_cet_util())
	def get_username(self) : return self.username
	def get_bm_util(self) : return self.bm_util
	def get_decompt_util(self) : return self.decompt_util
	def get_gpe_util(self) : return self.gpe_util
	def get_mess(self) : return self.mess
	def get_statut_util(self) : return self.statut_util
	def get_trans_cet_util(self) : return self.trans_cet_util
	def get_type_util(self) : return self.type_util
	def get_abs_set(self) : return self.tabsence_set
	def get_bm_util_set(self) : return self.tbonusmalusutilisateur_set
	def get_decompt_util_set(self) : return self.tdecomptesutilisateur_set
	def get_gpe_util_set(self) : return self.tgroupesutilisateur_set
	def get_mess_util_set(self) : return self.tmessagesutilisateur_set
	def get_role_util_set(self) : return self.trolesutilisateur_set
	def get_statut_util_set(self) : return self.tstatutsutilisateur_set
	def get_trans_cet_util_set(self) : return self.ttransactionscetutilisateur_set
	def get_verif_abs_set(self) : return self.tverificationabsence_set

	def __str__(self) : return self.get_username()

	def delete(self, *args, **kwargs) :

		# Désassignation de l'utilisateur en cours de suppression afin d'éviter la suppression en cascade des absences
		# dont il a juste émis ou vérifié
		TVerificationAbsence.objects.filter(id_util_verif = self.get_pk()).update(id_util_verif = None)
		TAbsence.objects.filter(id_util_connect = self.get_pk()).update(id_util_connect = None)
		
		super(TUtilisateur, self).delete(*args, **kwargs)

	'''
	Obtention des messages utilisateur
	self : Instance TUtilisateur
	_est_arch : "Boîte de réception" ou "Messages archivés" ?
	_lim : Limite du jeu de données
	Retourne un jeu de données
	'''
	def get_qs_mess_util_set(self, _est_arch, _lim = None) :
		qs = self.get_mess_util_set().filter(est_arch = _est_arch)
		if _lim :
			qs = qs[:_lim]
		return qs

	'''
	Obtention du nombre de messages utilisateur
	self : Instance TUtilisateur
	_est_arch : "Boîte de réception" ou "Messages archivés" ?
	Retourne un nombre entier
	'''
	def get_mess_util_set__count(self, _est_arch, _lim = None) :
		return self.get_mess_util_set().filter(est_arch = _est_arch).count()

	'''
	Obtention du nom complet d'un utilisateur
	self : Instance TUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_nom_complet(self) : return '{0} {1}'.format(self.get_last_name(), self.get_first_name())

	'''
	Obtention du statut utilisateur courant
	self : Instance TUtilisateur
	Retourne un nombre entier
	'''
	def get_statut_util_set__last(self) :
		
		# Obtention d'une instance TStatutsUtilisateur (la plus récente)
		obj_statut_util = self.get_statut_util_set().last()

		return obj_statut_util.get_statut_util__str() if obj_statut_util else None

	'''
	Obtention du nombre de messages non-lus de la boîte de réception
	self : Instance TUtilisateur
	Retourne un nombre entier
	'''
	def get_mess_util_set_bdr_non_lus__count(self) :
		return self.get_mess_util_set().filter(est_arch = False, est_lu = False).count()

	'''
	Obtention des rôles utilisateur sous forme de tableau
	self : Instance TUtilisateur
	Retourne un tableau
	'''
	def get_type_util__list(self) : return [tu.get_pk() for tu in self.get_type_util().all()]

	'''
	Mise à jour de la table t_decomptes_utilisateur après une requête d'insertion, de modification ou de suppression
	self : Instance TUtilisateur
	'''
	def calc_proratas(self) :

		# Imports
		from app.models import TAnnee
		from app.models import TDecomptesUtilisateur
		from django.conf import settings

		'''
		Obtention d'un décompte agent
		_tab : Tableau de données
		_code : RTT ou congés ?
		Retourne un nombre à une décimale au maximum
		'''
		def calc(_tab, _code) :

			# Initialisation des paramètres d'absence selon le statut agent
			if _code == 'R' :
				tab = { str(elem['us']) : elem['r_max'] for elem in settings.USER_STATUS }
			elif _code == 'C' :
				tab = { str(elem['us']) : elem['c_max'] for elem in settings.USER_STATUS }

			# Calcul du prorata
			prorata = 0
			for elem in _tab : prorata += round(2 * (tab[elem[0]] * elem[1] / 12)) / 2

			return prorata

		tab_decompt_util = []
		for a in TAnnee.objects.reverse() :

			# Initialisation des statuts agent pour l'année X
			tab_statut_util = [(
				su.get_statut_util(), su.get_mois_deb_statut_util()
			) for su in self.get_statut_util_set().filter(num_annee = a)]

			# Élimination potentielle d'un statut agent si début en janvier de l'année X
			trouve = False
			for elem in tab_statut_util :
				if elem[0] == 1 : trouve = True

			# Ajout (si existence) du dernier statut agent précédent l'année X
			if trouve == False :
				obj_statut_util = self.get_statut_util_set().filter(num_annee__lt = a).last()
				if obj_statut_util : tab_statut_util.insert(0, (obj_statut_util.get_statut_util(), 1))

			# Initialisation des proratas pour l'année X
			tab_donn = []
			for i in range(len(tab_statut_util)) :

				# Obtention de l'élément courant
				elem = tab_statut_util[i]

				# Calcul du nombre de mois travaillés
				try :
					nbre_mois = tab_statut_util[i + 1][1] - elem[1]
				except :
					nbre_mois = 13 - elem[1]

				tab_donn.append([str(elem[0]), nbre_mois])

			# Préparation des données à insérer dans la base de données
			if len(tab_donn) > 0 :
				tab_decompt_util.append({
					'nbre_conges_decompt_util' : calc(tab_donn, 'C'),
					'nbre_rtt_decompt_util' : calc(tab_donn, 'R'),
					'id_util' : self,
					'num_annee' : a
				})

		# Mise à jour des décomptes agent (suppression + ajout)
		self.get_decompt_util_set().all().delete()
		for elem in tab_decompt_util :
			TDecomptesUtilisateur.objects.create(
				nbre_conges_decompt_util = elem['nbre_conges_decompt_util'],
				nbre_rtt_decompt_util = elem['nbre_rtt_decompt_util'],
				id_util = elem['id_util'],
				num_annee = elem['num_annee']
			)

	'''
	Obtention des absences en attente pouvant être vérifiées par un utilisateur
	_self : Instance TUtilisateur
	_filter : Conditions de filtre
	Retourne un tableau
	'''
	def get_abs_a_verif__list(self, _filter = {}) :

		# Import
		from app.models import TAbsence

		# Initialisation des absences en attente pouvant être vérifiées par l'utilisateur
		output = []
		for a in TAbsence.objects.filter(**_filter) :
			if a.get_etat_abs() == 0 :
				if a.get_type_abs().get_gpe_type_abs().get_type_util().get_pk() in self.get_type_util__list() :
					output.append(a)

		return output

	'''
	Obtention des données utiles à la formation de la fiche récapitulative du CET
	self : Instance TUtilisateur
	_mt : Mode de tri (TN, TAC ou TAD)
	_lim : Limite en nombre d'années
	Retourne un tableau associatif
	'''
	def get_tabl_cet(self, _mt, _lim = None) :

		# Imports
		from app.functions import calcul_nbre_j
		from app.models import TVerificationAbsence
		from datetime import date
		from django.conf import settings
		import operator

		output = []

		# Initialisation du tableau de pré-sortie
		tab = []
		
		# Empilement de la sortie (transactions entrantes)
		for tcu in self.get_trans_cet_util_set().all() :
			tab.append({
				'annee' : tcu.get_annee().get_pk(),
				'dt' : tcu.get_dt_trans_cet_util(),
				'nbre_conges_trans_cet_util' : tcu.get_nbre_conges_trans_cet_util__str(), 
				'nbre_rtt_trans_cet_util' : tcu.get_nbre_rtt_trans_cet_util__str(),
				'pk' : tcu.get_pk(),
				'terme' : tcu.get_nbre_conges_trans_cet_util() + tcu.get_nbre_rtt_trans_cet_util()
			})

		# Empilement de la sortie (transactions sortantes)
		for va in TVerificationAbsence.objects.filter(
			est_autor = True,
			id_abs_mere__id_util_emett = self,
			id_type_abs_final__id_gpe_type_abs = settings.DB_PK_DATAS['CET_PK']
		) :

			nbre_j_abs = - calcul_nbre_j({
				'dt_abs' : va.get_abs().get_dt_abs(), 'indisp_abs' : va.get_abs().get_indisp_abs()
			})

			tab.append({
				'annee' : va.get_abs().get_annee().get_pk(),
				'dt' : va.get_abs().get_dt_emiss_abs(),
				'nbre_j_abs' : '{0:g}'.format(nbre_j_abs),
				'dt_abs' : va.get_abs().get_dt_abs__fr_str(),
				'terme' : nbre_j_abs
			})

		# Initialisation des paramètres d'affichage
		if _mt == 'TN' :
			tab_params = [['dt'], False]
		elif _mt == 'TAC' :
			tab_params = [['annee', 'dt'], False]
		elif _mt == 'TAD' :
			tab_params = [['annee', 'dt'], True]
		else :
			tab_params = None

		if tab_params :

			# Tri du tableau
			tab = sorted(tab, key = operator.itemgetter(*tab_params[0]))

			# Détermination pour chaque transaction du total en cours à partir du solde initial
			total = self.get_solde_cet_util()
			for elem in tab : total += elem['terme']; elem['total'] = '{0:g}'.format(total)

			# Je renverse le tableau si besoin
			if tab_params[1] == True : tab.reverse()

		else :
			tab = []

		# Préparation du tableau de sortie
		if _lim :
			for elem in tab :
				if elem['annee'] >= date.today().year - _lim : output.append(elem)
		else :
			output = tab

		return output

	'''
	Calcul du solde restant sur le CET
	self : Instance TUtilisateur
	_apercu : Nombre entier ou à une décimale simulant le solde initial
	Retourne un nombre entier ou à une décimale
	'''
	def get_solde_cet_restant_util(self, _apercu = None) :

		# Imports
		from app.functions import calcul_nbre_j
		from app.models import TVerificationAbsence
		from django.conf import settings

		# Initialisation du solde restant (= solde initial)
		output = _apercu if _apercu is not None else self.get_solde_cet_util()
		
		# Calcul du solde restant (= solde initial + entrées - sorties)
		for tcu in self.get_trans_cet_util_set().all() :
			output += (tcu.get_nbre_conges_trans_cet_util() + tcu.get_nbre_rtt_trans_cet_util())

		for va in TVerificationAbsence.objects.filter(
			est_autor = True,
			id_abs_mere__id_util_emett = self,
			id_type_abs_final__id_gpe_type_abs = settings.DB_PK_DATAS['CET_PK']
		) :
			output -= calcul_nbre_j({
				'dt_abs' : va.get_abs().get_dt_abs(), 'indisp_abs' : va.get_abs().get_indisp_abs()
			})

		return output

	'''
	Obtention du solde restant sur le CET sous forme de chaîne de caractères
	self : Instance TUtilisateur
	_apercu : Nombre entier ou à une décimale simulant le solde initial
	Retourne une chaîne de caractères
	'''
	def get_solde_cet_restant_util__str(self, _apercu = None) :
		return '{0:g}'.format(self.get_solde_cet_restant_util(_apercu))

	'''
	Obtention du nombre de jours de congés payés ou de RTT autorisés pour une année donnée
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne un nombre entier ou à une décimale
	'''
	def get_nbre_j_autor(self, _cp, _annee) :

		# Imports
		from app.functions import calcul_nbre_j
		from app.models import TVerificationAbsence
		from django.conf import settings

		output = 0

		# Initialisation du groupe de type d'absence
		gta_pk = settings.DB_PK_DATAS['C_PK'] if _cp == True else settings.DB_PK_DATAS['RTT_PK']

		# Calcul du nombre de jours autorisés
		for va in TVerificationAbsence.objects.filter(
			est_autor = True,
			id_abs_mere__id_util_emett = self,
			id_abs_mere__num_annee = _annee,
			id_type_abs_final__id_gpe_type_abs = gta_pk
		) :
			output += calcul_nbre_j({
				'dt_abs' : va.get_abs().get_dt_abs(), 'indisp_abs' : va.get_abs().get_indisp_abs()
			})

		return output

	'''
	Obtention du nombre de jours de congés payés ou de RTT autorisés pour une année donnée sous forme de chaîne de
	caractères
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_autor__str(self, _cp, _annee) : return '{0:g}'.format(self.get_nbre_j_autor(_cp, _annee))

	'''
	Obtention du nombre de jours de congés payés ou de RTT restants pour une année donnée
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_rest(self, _cp, _annee) :

		# Obtention d'une instance TDecomptesUtilisateur
		obj_decompt_util = self.get_decompt_util_set().get(num_annee = _annee)

		# Définition du nombre de jours de base
		if _cp == True :
			nbre_j_base = obj_decompt_util.get_nbre_j_cp_base()
		else :
			nbre_j_base = obj_decompt_util.get_nbre_j_rtt_base()

		return nbre_j_base - self.get_nbre_j_autor(_cp, _annee)

	'''
	Obtention du nombre de jours de congés payés ou de RTT restants pour une année donnée sous forme de chaîne de
	caractères
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_rest__str(self, _cp, _annee) : return '{0:g}'.format(self.get_nbre_j_rest(_cp, _annee))

	'''
	Obtention du nombre de jours de congés payés ou de RTT transférés sur le CET pour une année donnée
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne un nombre entier ou à une décimale
	'''
	def get_nbre_j_transf(self, _cp, _annee) :

		output = 0
		
		# Calcul du solde restant (= solde initial + entrées - sorties)
		for tcu in self.get_trans_cet_util_set().filter(num_annee = _annee) :
			output += tcu.get_nbre_conges_trans_cet_util() if _cp == True else tcu.get_nbre_rtt_trans_cet_util()

		return output

	'''
	Obtention du nombre de jours de congés payés ou de RTT transférés sur le CET pour une année donnée sous forme de
	chaîne de caractères
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_transf__str(self, _cp, _annee) : return '{0:g}'.format(self.get_nbre_j_transf(_cp, _annee))

	'''
	Obtention du nombre de jours de congés payés ou de RTT non-transférés sur le CET pour une année donnée
	self : Instance TUtilisateur
	_cp : Congés payés ?
	_annee : Année
	_tcu : Instance TTransactionsCetUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_non_transf(self, _cp, _annee, _tcu = None) :

		# Calcul du nombre de jours non-transférés sur le CET
		nbre_j_non_transf = self.get_nbre_j_rest(_cp, _annee) - self.get_nbre_j_transf(_cp, _annee)

		# Omission d'une instance TTransactionsCetUtilisateur en cas de mise à jour
		if _tcu :
			if _cp == True :
				nbre_j_non_transf += _tcu.get_nbre_conges_trans_cet_util()
			else :
				nbre_j_non_transf += _tcu.get_nbre_rtt_trans_cet_util()

		return nbre_j_non_transf

	'''
	Obtention d'une fiche CET au format Word
	self : Instance TUtilisateur
	_mt : Mode de tri (TN, TAC ou TAD)
	_lim : Limite en nombre d'années
	Retourne un document Word
	'''
	def get_fiche_cet(self, _mt, _lim = None) :

		# Imports
		from app.functions import gener_cdc
		from datetime import date
		from docx import Document
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		from docx.oxml import parse_xml
		from docx.oxml.ns import nsdecls
		from docx.shared import Pt
		from django.http import HttpResponse

		# Déclaration d'un nouveau document Word
		document = Document()

		# Définition des styles globaux
		style = document.styles['Normal']
		font = style.font
		font.name = 'Calibri'
		font.size = Pt(10)

		# Définition du titre du document et de ses styles
		titre = 'COMPTE ÉPARGNE TEMPS'
		annee = date.today().year
		if _lim : titre += ' (DE {0} À {1})'.format(annee - _lim, annee)
		paragraph = document.add_paragraph()
		run = paragraph.add_run(titre)
		run.font.size = Pt(20); run.bold = True; run.underline = True;
		paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Définition du titulaire de la fiche CET et de ses styles
		paragraph = document.add_paragraph()
		run = paragraph.add_run('Nom complet :')
		run.underline = True
		run = paragraph.add_run(' {}'.format(self.get_nom_complet()))

		# Déclaration du tableau des transactions sur le CET
		table = document.add_table(rows = 3, cols = 6, style = 'TableGrid')
		table.autofit = True

		# Définition du texte de chaque cellule de la ligne en-tête
		table.cell(0, 0).text = 'Année'
		table.cell(0, 1).text = 'Épargne'
		table.cell(1, 1).text = 'Entrées'
		table.cell(2, 1).text = 'Congés payés'
		table.cell(2, 2).text = 'RTT'
		table.cell(1, 3).text = 'Sorties'
		table.cell(0, 4).text = 'Total'
		table.cell(0, 5).text = 'Congés pris'

		# Fusion verticale
		table.cell(0, 1).merge(table.cell(0, 3))
		table.cell(1, 1).merge(table.cell(1, 2))

		# Fusion horizontale
		table.cell(0, 0).merge(table.cell(2, 0))
		table.cell(1, 3).merge(table.cell(2, 3))
		table.cell(0, 4).merge(table.cell(2, 4))
		table.cell(0, 5).merge(table.cell(2, 5))

		# Mise en forme générale de l'en-tête du tableau des transactions sur le CET
		for lg in table.rows :
			for cell in lg.cells :
				cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F6F6F6"/>'.format(nsdecls('w'))))
				cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		for elem in self.get_tabl_cet(_mt, _lim) :
			
			# Initialisation des données de la transaction courante
			tab_donn = [
				elem['annee'],
				elem['nbre_conges_trans_cet_util'] if 'nbre_conges_trans_cet_util' in elem else '',
				elem['nbre_rtt_trans_cet_util'] if 'nbre_rtt_trans_cet_util' in elem else '',
				elem['nbre_j_abs'] if 'nbre_j_abs' in elem else '',
				elem['total'],
				elem['dt_abs'] if 'dt_abs' in elem else ''
			]

			# Ajout d'une ligne
			row_cells = table.add_row().cells

			# Remplissage du tableau des transactions sur le CET
			for i in range(len(row_cells)) : row_cells[i].text = str(tab_donn[i])

		# Sauvegarde du document Word et téléchargement de celui-ci
		output = HttpResponse(content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		output['Content-Disposition'] = 'attachment; filename={}.docx'.format(gener_cdc())
		document.save(output)

		return output

	'''
	Obtention de la datatable relative aux absences autorisées ou en attente d'un utilisateur
	self : Instance TUtilisateur
	_annee : Nombre entier
	Retourne une chaîne de caractères
	'''
	def get_tabl_abs(self, _annee = None) :

		# Imports
		from app.models import TAbsence
		from django.core.urlresolvers import reverse

		# Initialisation du jeu de données des absences
		qs_abs = self.get_abs_set().filter(num_annee = _annee) if _annee else TAbsence.objects.none()

		# Initialisation des données de la datatable
		tab = []
		for a in qs_abs :
			if a.get_etat_abs() > -1 :
				tab.append([
					a.get_util_emett().get_nom_complet(),
					a.get_type_abs_final(),
					a.get_dt_abs__fr_str(),
					a.get_annee(),
					a.get_etat_abs(True),
					'<a href="{}" class="inform-icon pull-right" title="Consulter l\'absence"></a>'.format(
						reverse('consult_abs', args = [a.get_pk()])
					)
				])

		return '''
		<div class="custom-table" id="dtab_consult_abs">
			<table border="1" bordercolor="#DDD">
				<thead>
					<tr>
						<th>Nom complet de l'agent</th>
						<th>Type de l'absence</th>
						<th>Date de l'absence</th>
						<th>Année de l'absence</th>
						<th>État de l'absence</th>
						<th></th>
					</tr>
				</thead>
				<tbody>{}</tbody>
			</table>
		</div>
		'''.format(''.join(['<tr>{}</tr>'.format(''.join(['<td>{}</td>'.format(td) for td in tr])) for tr in tab]))

class TBonusMalusUtilisateur(models.Model) :

	# Imports
	from app.validators import valid_cdc
	from app.validators import valid_float

	# Attributs
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE, verbose_name = 'Agent concerné')
	num_annee = models.ForeignKey(TAnnee, on_delete = models.CASCADE, verbose_name = 'Année')
	int_bm_util = models.CharField(
		blank = True, max_length = 255, null = True, validators = [valid_cdc], verbose_name = 'Intitulé'
	)
	nbre_conges_bm_util = models.FloatField(
		default = 0, validators = [valid_float], verbose_name = 'Bonus/malus en congés payés'
	)
	nbre_rtt_bm_util = models.FloatField(default = 0, validators = [valid_float], verbose_name = 'Bonus/malus en RTT')

	class Meta :
		db_table = 't_bonus_malus_utilisateur'
		ordering = ['num_annee', 'id_util']
		verbose_name = verbose_name_plural = 'T_BONUS_MALUS_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_util(self) : return self.id_util
	def get_annee(self) : return self.num_annee
	def get_int_bm_util(self) : return self.int_bm_util
	def get_nbre_conges_bm_util(self) : return self.nbre_conges_bm_util
	def get_nbre_rtt_bm_util(self) : return self.nbre_rtt_bm_util

	def __str__(self) : return '{0} - {1}'.format(self.get_util(), self.get_annee())

class TDecomptesUtilisateur(models.Model) :

	# Attributs
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)
	num_annee = models.ForeignKey(TAnnee, on_delete = models.CASCADE)
	nbre_conges_decompt_util = models.FloatField()
	nbre_rtt_decompt_util = models.FloatField()

	class Meta :
		db_table = 't_decomptes_utilisateur'
		ordering = ['-num_annee']
		verbose_name = verbose_name_plural = 'T_DECOMPTES_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_util(self) : return self.id_util
	def get_annee(self) : return self.num_annee
	def get_nbre_conges_decompt_util(self) : return self.nbre_conges_decompt_util
	def get_nbre_conges_decompt_util__str(self) : return '{0:g}'.format(self.get_nbre_conges_decompt_util())
	def get_nbre_rtt_decompt_util(self) : return self.nbre_rtt_decompt_util
	def get_nbre_rtt_decompt_util__str(self) : return '{0:g}'.format(self.get_nbre_rtt_decompt_util())

	def __str__(self) : return '{0} - {1}'.format(self.get_util(), self.get_annee())

	'''
	Obtention du nombre de jours de congés payés de base avec prise en compte des bonus/malus
	self : Instance TDecomptesUtilisateur
	Retourne un nombre entier ou à une décimale
	'''
	def get_nbre_j_cp_base(self) :
		output = self.get_nbre_conges_decompt_util()
		for bmu in self.get_util().get_bm_util_set().filter(num_annee = self.get_annee()) :
			output += bmu.get_nbre_conges_bm_util()
		return output

	'''
	Obtention du nombre de jours de congés payés de base avec prise en compte des bonus/malus sous forme de chaîne de
	caractères
	self : Instance TDecomptesUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_cp_base__str(self) : return '{0:g}'.format(self.get_nbre_j_cp_base())

	'''
	Obtention du nombre de jours de RTT de base avec prise en compte des bonus/malus
	self : Instance TDecomptesUtilisateur
	Retourne un nombre entier ou à une décimale
	'''
	def get_nbre_j_rtt_base(self) :
		output = self.get_nbre_rtt_decompt_util()
		for bmu in self.get_util().get_bm_util_set().filter(num_annee = self.get_annee()) :
			output += bmu.get_nbre_rtt_bm_util()
		return output

	'''
	Obtention du nombre de jours de RTT de base avec prise en compte des bonus/malus sous forme de chaîne de caractères
	self : Instance TDecomptesUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_nbre_j_rtt_base__str(self) : return '{0:g}'.format(self.get_nbre_j_rtt_base())

class TGroupesUtilisateur(models.Model) :

	# Attributs
	id_gpe_util = models.ForeignKey(TGroupeUtilisateur, on_delete = models.CASCADE)
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)

	class Meta :
		db_table = 't_groupes_utilisateur'
		verbose_name = verbose_name_plural = 'T_GROUPES_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_gpe_util(self) : return self.id_gpe_util
	def get_util(self) : return self.id_util

	def __str__(self) : return '{0} - {1}'.format(self.get_gpe_util(), self.get_util())

class TMessagesUtilisateur(models.Model) :

	id_mess = models.ForeignKey(TMessage, on_delete = models.CASCADE)
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)
	est_arch = models.BooleanField(default = False)
	est_lu = models.BooleanField(default = False)

	class Meta :
		db_table = 't_messages_utilisateur'
		ordering = ['id_mess']
		verbose_name = verbose_name_plural = 'T_MESSAGES_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_mess(self) : return self.id_mess
	def get_util(self) : return self.id_util
	def get_est_arch(self) : return self.est_arch
	def get_est_lu(self) : return self.est_lu

	def __str__(self) : return '{0} - {1}'.format(self.get_mess(), self.get_util())

	def delete(self, *args, **kwargs) :

		# Obtention d'une instance TMessage
		obj_mess = self.get_mess()

		# Suppression d'une instance TMessagesUtilisateur
		super(TMessagesUtilisateur, self).delete(*args, **kwargs)

		# Suppression d'une instance TMessage si aucun raccordement avec une autre instance TMessagesUtilisateur
		if TMessagesUtilisateur.objects.filter(id_mess = obj_mess).count() == 0 : obj_mess.delete()

class TRolesUtilisateur(models.Model) :

	# Attributs
	code_type_util = models.ForeignKey(TTypeUtilisateur, on_delete = models.DO_NOTHING)
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)

	class Meta :
		db_table = 't_roles_utilisateur'
		verbose_name = verbose_name_plural = 'T_ROLES_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_type_util(self) : return self.code_type_util
	def get_util(self) : return self.id_util

	def __str__(self) : return '{0} - {1}'.format(self.get_type_util(), self.get_util())

class TStatutsUtilisateur(models.Model) :

	# Imports
	from app.functions import get_obj_dt
	from django.conf import settings

	# Attributs
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)
	num_annee = models.ForeignKey(TAnnee, on_delete = models.CASCADE, verbose_name = 'Année de mise en place')
	mois_deb_statut_util = models.IntegerField(
		choices = [(index + 1, elem) for index, elem in enumerate(get_obj_dt('MONTHS'))],
		verbose_name = 'Mois de mise en place'
	)
	statut_util = models.IntegerField(
		choices = [(elem['us'], '{} %'.format(elem['us'])) for elem in settings.USER_STATUS],
		verbose_name = 'Temps de travail hebdomadaire'
	)

	class Meta :
		db_table = 't_statuts_utilisateur'
		ordering = ['-num_annee', 'mois_deb_statut_util']
		verbose_name = verbose_name_plural = 'T_STATUTS_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_util(self) : return self.id_util
	def get_annee(self) : return self.num_annee
	def get_mois_deb_statut_util(self) : return self.mois_deb_statut_util
	def get_statut_util(self) : return self.statut_util
	def get_statut_util__str(self) : return '{} %'.format(self.get_statut_util())

	def __str__(self) : return '{0} - {1}'.format(self.get_util(), self.get_annee())

	def delete(self, *args, **kwargs) :

		# Obtention d'une instance TUtilisateur
		obj_util = self.get_util()

		# Suppression d'une instance TStatutsUtilisateur
		super(TStatutsUtilisateur, self).delete(*args, **kwargs)

		# Mise à jour de la table t_decomptes_utilisateur
		obj_util.calc_proratas()

	'''
	Obtention du mois et de l'année de début d'un statut utilisateur
	self : Instance TStatutsUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_period_deb_statut_util(self) :

		# Import
		from app.functions import get_obj_dt

		return '{0} {1}'.format(get_obj_dt('MONTHS')[self.get_mois_deb_statut_util() - 1], self.get_annee())

	'''
	Obtention de la date de début d'un statut utilisateur
	self : Instance TStatutsUtilisateur
	Retourne un objet "date"
	'''
	def get_dt_deb_statut_util(self) :
		from datetime import date; return date(self.get_annee().get_pk(), self.get_mois_deb_statut_util(), 1)

	'''
	Obtention de la date de début d'un statut utilisateur
	self : Instance TStatutsUtilisateur
	Retourne une chaîne de caractères
	'''
	def get_dt_deb_statut_util__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_dt_deb_statut_util())

class TTransactionsCetUtilisateur(models.Model) :

	# Import
	from app.validators import valid_float

	# Attributs
	id_util = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)
	num_annee = models.ForeignKey(TAnnee, on_delete = models.CASCADE)
	dt_trans_cet_util = models.DateTimeField()
	nbre_conges_trans_cet_util = models.FloatField(
		default = 0, validators = [valid_float], verbose_name = 'Nombre de jours de congés à transférer sur le CET'
	)
	nbre_rtt_trans_cet_util = models.FloatField(
		default = 0, validators = [valid_float], verbose_name = 'Nombre de jours de RTT à transférer sur le CET'
	)

	class Meta :
		db_table = 't_transactions_cet_utilisateur'
		verbose_name = verbose_name_plural = 'T_TRANSACTIONS_CET_UTILISATEUR'

	# Getters
	def get_pk(self) : return self.pk
	def get_util(self) : return self.id_util
	def get_annee(self) : return self.num_annee
	def get_dt_trans_cet_util(self) : return self.dt_trans_cet_util
	def get_dt_trans_cet_util__str(self) :
		from app.functions import get_local_format
		return get_local_format(self.get_dt_trans_cet_util())
	def get_nbre_conges_trans_cet_util(self) : return self.nbre_conges_trans_cet_util
	def get_nbre_conges_trans_cet_util__str(self) : return '{0:g}'.format(self.get_nbre_conges_trans_cet_util())
	def get_nbre_rtt_trans_cet_util(self) : return self.nbre_rtt_trans_cet_util
	def get_nbre_rtt_trans_cet_util__str(self) : return '{0:g}'.format(self.get_nbre_rtt_trans_cet_util())

	def __str__(self) : return '{0} - {1}'.format(self.get_util(), self.get_annee())

class TGroupeTypeAbsence(models.Model) :

	# Import
	from app.validators import valid_cdc

	# Attributs
	id_gpe_type_abs = models.AutoField(primary_key = True)
	abrev_gpe_type_abs = models.CharField(max_length = 255, verbose_name = 'Abréviation')
	coul_gpe_type_abs = models.CharField(
		default = 'revert',
		help_text = '''
		Une couleur valide doit respecter le code couleur HTML. Afin de vous aider à choisir une couleur valide,
		veuillez cliquer <a href="http://www.toutes-les-couleurs.com/code-couleur-html.php" target="blank">ici</a>.
		''',
		max_length = 255,
		validators = [valid_cdc],
		verbose_name = 'Couleur utilisée dans le calendrier des absences'
	)
	est_disp = models.BooleanField(verbose_name = 'Est sélectionnable par les agents ou la direction ?')
	int_gpe_type_abs = models.CharField(max_length = 255, validators = [valid_cdc], verbose_name = 'Intitulé')
	ordre_zl_gpe_type_abs = models.IntegerField(
		blank = True, null = True, verbose_name = 'Ordre dans la liste déroulante'
	)
	code_type_util = models.ForeignKey(
		TTypeUtilisateur,
		on_delete = models.DO_NOTHING,
		verbose_name = 'Type d\'utilisateur éligible à la vérification'
	)

	class Meta :
		db_table = 't_groupe_type_absence'
		ordering = ['ordre_zl_gpe_type_abs', 'int_gpe_type_abs']
		verbose_name = verbose_name_plural = 'T_GROUPE_TYPE_ABSENCE'

	# Getters
	def get_pk(self) : return self.pk
	def get_abrev_gpe_type_abs(self) : return self.abrev_gpe_type_abs
	def get_coul_gpe_type_abs(self) : return self.coul_gpe_type_abs
	def get_est_disp(self) : return self.est_disp
	def get_int_gpe_type_abs(self) : return self.int_gpe_type_abs
	def get_ordre_zl_gpe_type_abs(self) : return self.ordre_zl_gpe_type_abs
	def get_type_util(self) : return self.code_type_util
	def get_type_abs_set(self) : return self.ttypeabsence_set

	def __str__(self) : return self.get_int_gpe_type_abs()

class TTypeAbsence(models.Model) :

	# Import
	from app.validators import valid_cdc

	# Attributs
	id_type_abs = models.AutoField(primary_key = True)
	descr_nds_type_abs = models.TextField(
		blank = True,
		null = True,
		validators = [valid_cdc],
		verbose_name = 'Description de la note de service'
	)
	int_type_abs = models.CharField(max_length = 255, validators = [valid_cdc], verbose_name = 'Intitulé')
	pj_abs_req = models.BooleanField(verbose_name = 'La pièce justificative est-elle obligatoire ?')
	id_gpe_type_abs = models.ForeignKey(
		TGroupeTypeAbsence, on_delete = models.CASCADE, verbose_name = 'Groupe de type d\'absence'
	)

	class Meta :
		db_table = 't_type_absence'
		ordering = ['int_type_abs']
		verbose_name = verbose_name_plural = 'T_TYPE_ABSENCE'

	# Getters
	def get_pk(self) : return self.pk
	def get_descr_nds_type_abs(self) : return self.descr_nds_type_abs
	def get_int_type_abs(self) : return self.int_type_abs
	def get_pj_abs_req(self) : return self.pj_abs_req
	def get_gpe_type_abs(self) : return self.id_gpe_type_abs
	def get_abs_set(self) : return self.tabsence_set

	def __str__(self) : return self.get_int_type_abs()

class TAbsence(models.Model) :

	# Imports
	from app.validators import valid_cdc
	from app.validators import valid_pdf
	from django.contrib.postgres.fields import ArrayField

	'''
	Définition du chemin d'upload du justificatif d'absence
	_inst : Instance TAbsence
	_fname : Nom du fichier
	Retourne une chaîne de caractères
	'''
	def set_pj_abs__upload_to(_inst, _fname) :
		from app.functions import gener_cdc; return 'justificatifs/{0}.{1}'.format(gener_cdc(), _fname.split('.')[-1])

	# Attributs
	id_abs = models.AutoField(primary_key = True)
	comm_abs = models.TextField(blank = True, null = True, validators = [valid_cdc], verbose_name = 'Commentaire')
	dt_abs = ArrayField(base_field = models.DateField(), size = 2)
	dt_emiss_abs = models.DateTimeField()
	indisp_abs = ArrayField(base_field = models.CharField(max_length = 2), size = 2)
	pj_abs = models.FileField(
		blank = True,
		null = True,
		upload_to = set_pj_abs__upload_to,
		validators = [valid_pdf],
		verbose_name = 'Insérer le justificatif d\'absence <span class="fl-complement">(fichier PDF)</span>'
	)
	src_alerte = models.BooleanField(verbose_name = 'Peut être soumise à une alerte ?')
	id_type_abs = models.ForeignKey(TTypeAbsence, on_delete = models.DO_NOTHING)
	id_util_connect = models.ForeignKey(TUtilisateur, on_delete = models.DO_NOTHING, null = True, related_name = '+')
	id_util_emett = models.ForeignKey(TUtilisateur, on_delete = models.CASCADE)
	num_annee = models.ForeignKey(TAnnee, on_delete = models.CASCADE)

	class Meta :
		db_table = 't_absence'
		ordering = ['-dt_abs', 'id_util_emett']
		verbose_name = verbose_name_plural = 'T_ABSENCE'

	# Getters
	def get_pk(self) : return self.pk
	def get_comm_abs(self) : return self.comm_abs
	def get_dt_abs(self) : return self.dt_abs
	def get_dt_abs__str(self) :
		from app.functions import get_local_format; return [get_local_format(elem) for elem in self.get_dt_abs()]
	def get_dt_emiss_abs(self) : return self.dt_emiss_abs
	def get_dt_emiss_abs__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_dt_emiss_abs())
	def get_indisp_abs(self) : return self.indisp_abs
	def get_pj_abs(self) : return self.pj_abs
	def get_pj_abs__path(self) :
		from django.conf import settings
		return settings.MEDIA_URL + self.get_pj_abs().name if self.get_pj_abs() else None
	def get_src_alerte(self) : return self.src_alerte
	def get_type_abs(self) : return self.id_type_abs
	def get_util_connect(self) : return self.id_util_connect
	def get_util_emett(self) : return self.id_util_emett
	def get_annee(self) : return self.num_annee
	def get_dt_abs_set(self) : return self.tdatesabsence_set
	def get_verif_abs(self) : return self.tverificationabsence if hasattr(self, 'tverificationabsence') else None

	def __str__(self) : return '{0} - {1} - {2}'.format(self.get_util_emett(), self.get_annee(), self.get_type_abs())

	def delete(self, *args, **kwargs) : self.get_pj_abs().delete(); super(TAbsence, self).delete(*args, **kwargs)
		
	'''
	Obtention de l'état de vérification d'une absence
	self : Instance TAbsence
	Retourne une chaîne de caractères
	'''
	def get_etat_abs(self, _str = False) :
		if self.get_verif_abs() :
			if self.get_verif_abs().get_est_autor() == True :
				output = { 'int' : 1, 'str' : 'Autorisée' }
			else :
				output = { 'int' : -1, 'str' : 'Refusée' }
		else :
			output = { 'int' : 0, 'str' : 'En attente' }
		return output['int'] if _str == False else output['str']

	'''
	Obtention d'une date avec précision
	self : Instance TAbsence
	Retourne une chaîne de caractères
	'''
	def get_dt_abs__fr_str(self) :

		'''
		Conversion d'un code d'indisponibilté
		_code : Code
		Retourne une chaîne de caractères
		'''
		def conv_indisp_abs(_code) :
			if _code == 'AM' :
				output = 'matin'
			elif _code == 'PM' :
				output = 'après-midi'
			elif _code == 'WD' :
				output = 'journée entière'
			return output

		# Initialisation de la sortie
		output = []
		for i in range(len(self.get_dt_abs__str())) :
			output.append('{0} ({1})'.format(self.get_dt_abs__str()[i], conv_indisp_abs(self.get_indisp_abs()[i])))

		return ' - '.join(output)

	'''
	Obtention du type d'absence final dans tous les cas de vérification
	self : Instance TAbsence
	Retourne une chaîne de caractères
	'''
	def get_type_abs_final(self) :

		# Détermination du type d'absence final dans tous les cas de vérification
		if self.get_verif_abs() :
			if self.get_verif_abs().get_est_autor() == True :
				output = self.get_verif_abs().get_type_abs_final()
			else :
				output = self.get_type_abs()
		else :
			output = self.get_type_abs()

		return output

	'''
	Obtention du nombre de jours d'absence
	self : Instance TAbsence
	Retourne un nombre entier ou à une décimale
	'''
	def get_nbre_dt_abs(self) :

		output = 0

		# Calcul du nombre de jours d'absence
		for da in self.get_dt_abs_set().all() :
			if da.get_indisp_dt_abs() == 'WD' :
				output += 1
			else :
				output += 0.5

		return output

	'''
	Vérification du droit de lecture
	self : Instance TAbsence
	_u : Instance TUtilisateur
	_func : Fonction ou procédure ?
	Retourne un booléen si fonction
	'''
	def can_read(self, _u, _func) :

		# Import
		from django.core.exceptions import PermissionDenied

		if _func == False :
			if self.get_util_emett() != _u and 'S' not in _u.get_type_util__list() : raise PermissionDenied
		else :
			return False if self.get_util_emett() != _u and 'S' not in _u.get_type_util__list() else True

	def can_update_pj_abs(self, _u, _func) :

		# Import
		from django.core.exceptions import PermissionDenied

		autorise = True
		if self.can_read(_u, True) == False :
			autorise == False
		else :
			if 'S' not in _u.get_type_util__list() and self.get_verif_abs() : autorise = False

		# Définition de la sortie
		if _func == False :
			if autorise == False : raise PermissionDenied
		else :
			return autorise

	'''
	Vérification du droit de modification du type final d'une absence
	self : Instance TAbsence
	_u : Instance TUtilisateur
	_func : Fonction ou procédure ?
	Retourne un booléen si fonction
	'''
	def can_update_type_abs_final(self, _u, _func) :

		# Import
		from django.core.exceptions import PermissionDenied

		# Instance TVerificationAbsence
		o_verif_abs = self.get_verif_abs()

		if o_verif_abs and o_verif_abs.get_est_autor() == True and 'S' in _u.get_type_util__list() :
			peut = True
		else :
			peut = False

		# Définition de la sortie
		if _func == False :
			if peut == False : raise PermissionDenied
		else :
			return peut

	'''
	Vérification du droit de suppression
	self : Instance TAbsence
	_u : Instance TUtilisateur
	_func : Fonction ou procédure ?
	Retourne un booléen si fonction
	'''
	def can_delete(self, _u, _func) :

		# Imports
		from datetime import date
		from django.core.exceptions import PermissionDenied

		autorise = True
		if self.can_read(_u, True) == False :
			autorise == False
		else :
			if 'S' not in _u.get_type_util__list() and date.today() >= self.get_dt_abs()[0] : autorise = False

		# Définition de la sortie
		if _func == False :
			if autorise == False : raise PermissionDenied
		else :
			return autorise

class TDatesAbsence(models.Model) :

	# Attributs
	id_dt_abs = models.AutoField(primary_key = True)
	dt_abs = models.DateField()
	indisp_dt_abs = models.CharField(max_length = 2)
	id_abs = models.ForeignKey(TAbsence, on_delete = models.CASCADE)

	class Meta :
		db_table = 't_dates_absence'
		ordering = ['dt_abs', 'indisp_dt_abs']
		verbose_name = verbose_name_plural = 'T_DATES_ABSENCE'

	# Getters
	def get_pk(self) : return self.pk
	def get_dt_abs(self) : return self.dt_abs
	def get_dt_abs__str(self) : from app.functions import get_local_format; return get_local_format(self.get_dt_abs())
	def get_indisp_dt_abs(self) : return self.indisp_dt_abs
	def get_abs(self) : return self.id_abs

	def __str__(self) : return self.get_dt_abs__str()

	'''
	Obtention de données utiles à la formation du calendrier des absences
	self : Instance TDatesAbsence
	Retourne un tableau associatif
	'''
	def get_donn_cal_abs(self) :

		# Instance TVerificationAbsence
		obj_verif_abs = self.get_abs().get_verif_abs()

		# Détermination du groupe de type d'absence
		if obj_verif_abs :
			obj_gpe_type_abs = obj_verif_abs.get_type_abs_final().get_gpe_type_abs()
		else :
			obj_gpe_type_abs = self.get_abs().get_type_abs_final().get_gpe_type_abs()

		return {
			'abrev_gpe_type_abs': obj_gpe_type_abs.get_abrev_gpe_type_abs(),
			'coul_gpe_type_abs' : obj_gpe_type_abs.get_coul_gpe_type_abs(),
			'int_dt_abs' : '{0} {1}'.format(
				obj_gpe_type_abs.get_abrev_gpe_type_abs(),
				{ 'AM' : '1/2 M', 'PM' : '1/2 A', 'WD' : 'JE' }[self.get_indisp_dt_abs()]
			)
		}

class TVerificationAbsence(models.Model) :

	# Import
	from app.validators import valid_cdc

	# Attributs
	id_abs_mere = models.OneToOneField(TAbsence, on_delete = models.CASCADE, primary_key = True)
	comm_verif_abs = models.TextField(
		blank = True, null = True, validators = [valid_cdc], verbose_name = 'Commentaire'
	)
	dt_verif_abs = models.DateTimeField()
	est_autor = models.BooleanField(default = False, verbose_name = 'Autorisez-vous l\'absence ?')
	id_type_abs_final = models.ForeignKey(TTypeAbsence, on_delete = models.DO_NOTHING)
	id_util_verif = models.ForeignKey(TUtilisateur, null = True, on_delete = models.DO_NOTHING)

	class Meta :
		db_table = 't_verification_absence'
		verbose_name = verbose_name_plural = 'T_VERIFICATION_ABSENCE'

	# Getters
	def get_pk(self) : return self.pk
	def get_comm_verif_abs(self) : return self.comm_verif_abs
	def get_dt_verif_abs(self) : return self.dt_verif_abs
	def get_dt_verif_abs__str(self) :
		from app.functions import get_local_format; return get_local_format(self.get_dt_verif_abs())
	def get_est_autor(self) : return self.est_autor
	def get_abs(self) : return self.id_abs_mere
	def get_type_abs_final(self) : return self.id_type_abs_final if self.get_est_autor() == True else None
	def get_util_verif(self) : return self.id_util_verif
	def get_dt_abs_set(self) : return self.tdatesabsence_set

	def __str__(self) : return str(self.get_abs())

class TDateFermeture(models.Model) :

	# Attribut
	id_dt_ferm = models.DateField(primary_key = True, verbose_name = 'Date de fermeture exceptionnelle')

	class Meta :
		db_table = 't_date_fermeture'
		ordering = ['pk']
		verbose_name = verbose_name_plural = 'T_DATE_FERMETURE'

	def __str__(self) : return self.get_pk__str()

	# Getters
	def get_pk(self) : return self.pk
	def get_pk__str(self) : from app.functions import get_local_format; return get_local_format(self.get_pk())